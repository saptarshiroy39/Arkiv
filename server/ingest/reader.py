import io
import csv
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import google.generativeai as genai
from server.config import genai as def_genai, logger

def read_pdf(blob, fname, api_key=None, **kwargs):
    try:
        if api_key:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
        else:
            model = def_genai.GenerativeModel("gemini-2.0-flash")
        
        pdf_part = {"mime_type": "application/pdf", "data": blob}
        prompt = "Extract all text content from this PDF document. Preserve the structure and formatting as much as possible. Include all text, tables, and data."
        res = model.generate_content([prompt, pdf_part])
        return [{"text": f"[Content from {fname}]\n{res.text}", "page": 1}]
    except Exception as e:
        logger.error(f"PDF error {fname}: {e}")
        return []


def read_docx(blob, fname, **kwargs):
    parts = []
    try:
        doc = Document(io.BytesIO(blob))
        for p in doc.paragraphs:
            if p.text.strip(): parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_txt = "\t".join([c.text for c in row.cells])
                if row_txt.strip(): parts.append(row_txt)
        return [{"text": "\n".join(parts), "page": 1}] if parts else []
    except Exception as e:
        logger.error(f"Docx error {fname}: {e}")
        return []


def read_excel(blob, fname, **kwargs):
    parts = []
    try:
        wb = load_workbook(io.BytesIO(blob), data_only=True)
        for sheet in wb:
            parts.append(f"--- {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                vals = [str(v) if v is not None else "" for v in row]
                line = "\t".join(vals)
                if line.strip(): parts.append(line)
        return [{"text": "\n".join(parts), "page": 1}] if parts else []
    except Exception as e:
        logger.error(f"Excel error {fname}: {e}")
        return []


def read_csv(blob, fname, **kwargs):
    try:
        txt = blob.decode('utf-8', errors='replace')
        content = "\n".join(["\t".join(r) for r in csv.reader(io.StringIO(txt)) if any(r)])
        return [{"text": content, "page": 1}] if content else []
    except Exception as e:
        logger.error(f"CSV error {fname}: {e}")
        return []


def read_pptx(blob, fname, api_key=None, **kwargs):
    try:
        if api_key:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
        else:
            model = def_genai.GenerativeModel("gemini-2.0-flash")
        
        pptx_part = {"mime_type": "application/vnd.openxmlformats-officedocument.presentationml.presentation", "data": blob}
        prompt = "Extract all text content from this PowerPoint presentation. For each slide, describe any charts, images, or diagrams. Preserve the structure and list slide numbers."
        res = model.generate_content([prompt, pptx_part])
        return [{"text": f"[Content from {fname}]\n{res.text}", "page": 1}]
    except Exception as e:
        logger.error(f"PPTX error {fname}: {e}")
        return []


def read_text(blob, fname, **kwargs):
    try:
        txt = blob.decode('utf-8', errors='ignore')
        return [{"text": txt, "page": 1}] if txt else []
    except Exception:
        return []


def read_image(blob, fname, api_key=None, **kwargs):
    try:
        if api_key:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
        else:
            model = def_genai.GenerativeModel("gemini-2.0-flash")

        img = Image.open(io.BytesIO(blob))
        prompt = "Describe this image in detail. Extract any text, charts, or key data point you see."
        
        res = model.generate_content([prompt, img])
        return [{"text": f"[Analysis for {fname}]\n{res.text}", "page": 1}]
    except Exception as e:
        logger.error(f"Image error {fname}: {e}")
        return []


def read_file(blob, fname, ftype, api_key=None):
    if ftype == "pdf": return read_pdf(blob, fname, api_key=api_key)
    if ftype == "docs": return read_docx(blob, fname)
    if ftype == "sheets": return read_excel(blob, fname)
    if ftype == "csv": return read_csv(blob, fname)
    if ftype == "slides": return read_pptx(blob, fname, api_key=api_key)
    if ftype == "image": return read_image(blob, fname, api_key=api_key)
    return read_text(blob, fname)
